from output import *
import re
import docx2txt
from docx import Document
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm,Pt
from docx.oxml.ns import qn
from operator import itemgetter
from image import *
import os
import shutil
import zipfile
import fnmatch

def replace_text(filename,old,new):
    doc = Document(filename)
    for p in doc.paragraphs:
        if old in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if old in inline[i].text:
                    text = inline[i].text.replace(old, new)
                    inline[i].text = text
    doc.save('test.docx')
    return
def replace_word():
    replace_text('./输出文档 - 对接1.docx','（园区名称）',park.loc['园区名称'][1])

    replace_text('test.docx','description',park.loc['园区描述'][1])

    replace_text('test.docx','城市气候描述',city.loc['城市气候描述'][1])

    replace_text('test.docx','园区规范',park.loc['园区规范'][1])

    replace_text('test.docx','土地使用情况',park.loc['土地使用情况'][1])

    replace_text('test.docx','位置描述',park.loc['位置描述'][1])

    replace_text('test.docx','城市名称',city.loc['城市名称'][1])

    replace_text('test.docx','平均温度',str(city.loc['平均温度'][1]))

    replace_text('test.docx','年最高温度',str(city.loc['年最高温度'][1]))

    replace_text('test.docx','年最低温度',str(city.loc['年最低温度'][1]))

    replace_text('test.docx','气候描述',city.loc['城市气候描述'][1])

    replace_text('test.docx','气候分区',city.loc['气候分区'][1])

    replace_text('test.docx','光伏情况',city.loc['光伏情况'][1])

    replace_text('test.docx','采暖供冷描述',city.loc['采暖供冷期描述'][1])

    replace_text('test.docx','电价描述',park.loc['电价描述'][1])

    replace_text('test.docx','用能政策',park.loc['用能政策'][1])

    replace_text('test.docx','地热资源评价',park.loc['地热资源评价'][1])

    replace_text('test.docx','卖电许可',park.loc['卖电许可'][1])

    replace_text('test.docx','氢价',str(park.loc['氢价'][1]))

    replace_text('test.docx','，以及卖电收益产生的抵扣。','。')

    replace_text('test.docx','供能对象描述',park.loc['供能对象描述'][1])

    replace_text('test.docx','供能方案描述',park.loc['用能方案描述'][1])

    replace_text('test.docx','所在省份',city.loc['所在省份'][1])

    replace_text('test.docx','供能方案描述',park.loc['用能方案描述'][1])

    replace_text('test.docx','制氢潜力',park.loc['制氢潜力'][1])

    replace_text('test.docx','（电负荷峰值数据）',str(grid_planning_output_json['ele_load_max']))

    replace_text('test.docx','（热负荷峰值数据）',str(grid_planning_output_json['g_demand_max']))

    replace_text('test.docx','（冷负荷峰值数据）',str(grid_planning_output_json['q_demand_max']))

    replace_text('test.docx','（电负荷总量数据）',str(grid_planning_output_json['ele_load_sum']))

    replace_text('test.docx','（热负荷总量数据）',str(grid_planning_output_json['g_demand_sum']))

    replace_text('test.docx','（冷负荷总量数据）',str(grid_planning_output_json['q_demand_sum']))

    replace_text('test.docx',"aaa",grid_planning_output_json['equipment_cost'])
    replace_text('test.docx','bbb',grid_operation_output_json['cer'])
    replace_text('test.docx','ccc',grid_operation_output_json['cer_perm2'])

    replace_text('test.docx','ia',itgrid_planning_output_json['equipment_cost'])
    replace_text('test.docx','薅',itgrid_operation_output_json['cer'])
    replace_text('test.docx','ic',itgrid_operation_output_json['cer_perm2'])

    replace_text('test.docx','eee',grid_planning_output_json['receive_year'])
    replace_text('test.docx','ite',itgrid_planning_output_json['receive_year'])

    replace_text('test.docx','fff',grid_operation_output_json['cer_rate'])
    replace_text('test.docx','if',itgrid_operation_output_json['cer_rate'])


def change_table_air_para(doc):
    tables = doc.tables[0]
    for i in range(len(tables.rows)):
        # 在第该表格i行列的单元格内输入对应内容”
        run = tables.cell(i, 1).paragraphs[0].add_run(str(list(gongnuan[1])[i]))
        # 设置字体
        run.font.name = 'Times New Roman'
        # 字体大小
        run.font.size = Pt(12)
        r = run._element
        # 中文宋体
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 居中
        tables.cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def change_table_equip_para(doc):
    tables = doc.tables[1]
    for i in range(1, len(tables.rows)):
        if str(list(equip['参数值'])[i - 1]) != 'nan':
            # 在第该表格i行2列的单元格内输入对应内容”
            run = tables.cell(i, 2).paragraphs[0].add_run(str(list(equip['参数值'])[i - 1]))
            # 设置字体
            run.font.name = 'Times New Roman'
            # 字体大小
            run.font.size = Pt(10.5)
            r = run._element
            # 中文宋体
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 居中
            tables.cell(i, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 在第该表格i行6列的单元格内输入对应内容”
        if str(list(equip['参数值'])[i + len(tables.rows) - 2]) != 'nan':
            run = tables.cell(i, 6).paragraphs[0].add_run(str(list(equip['参数值'])[i + len(tables.rows) - 2]))
            # 设置字体
            run.font.name = 'Times New Roman'
            # 字体大小
            run.font.size = Pt(10.5)
            r = run._element
            # 中文宋体
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 居中
            tables.cell(i, 6).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def change_table_equip_allocation(tables, json):
    keys = ["num_gtw",
            "p_fc_max",
            "p_hpg_max",
            "p_hp_max",
            "p_eb_max",
            "p_el_max",
            "hst",
            "m_ht",
            "m_ct",
            "area_pv",
            "area_sc",
            "p_co"]
    replace = itemgetter(*keys)(json)
    for i in range(1,len(tables.rows)):
        if replace[i-1] !=0 and replace[i-1] != '0.00':
            # 在第该表格i行列的单元格内输入对应内容”
            run = tables.cell(i, 2).paragraphs[0].add_run(str(replace[i-1]))
            # 设置字体
            run.font.name = 'Times New Roman'
            # 字体大小
            run.font.size = Pt(12)
            r = run._element
            # 中文宋体
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 居中
            tables.cell(i, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def change_table_eco_analyse(tables, planing_json,opertion_json):
    replace = [planing_json['equipment_cost'],
               opertion_json['operation_cost'],
               opertion_json['cost_save_rate'],
               planing_json['receive_year'],
               opertion_json['co2'],
               opertion_json['cer_rate']]
    for i in range(1,len(tables.rows)):
        run = tables.cell(i, 1).paragraphs[0].add_run(str(replace[i - 1]))
        # 设置字体
        run.font.name = 'Times New Roman'
        # 字体大小
        run.font.size = Pt(12)
        r = run._element
        # 中文宋体
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 居中
        tables.cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
def zip_docx(startdir, file_news):
    z = zipfile.ZipFile(file_news, 'w', zipfile.ZIP_DEFLATED)
    for dirpath, dirnames, filenames in os.walk(startdir):
        fpath = dirpath.replace(startdir, '')
        fpath = fpath and fpath + os.sep or ''
        for filename in filenames:
            z.write(os.path.join(dirpath, filename), fpath+filename)

if __name__=='__main__':
    ###读配置文件
    park = pd.read_excel('./配置文档数据.xlsx', sheet_name='园区文字', header=None, index_col=0)
    city = pd.read_excel('./配置文档数据.xlsx', sheet_name='城市文字', header=None, index_col=0)
    gongnuan = pd.read_excel('./配置文档数据.xlsx', sheet_name='供暖规范', header=None, index_col=0)
    equip = pd.read_excel('./配置文档数据.xlsx', sheet_name='设备参数', index_col=0)

    ###更改文字与表格
    replace_word()
    doc = Document('test.docx')

    change_table_air_para(doc)
    change_table_equip_para(doc)

    table2=doc.tables[2]
    table3 = doc.tables[3]
    table4 = doc.tables[4]
    table5 = doc.tables[5]

    change_table_equip_allocation(table2,grid_planning_output_json)
    change_table_equip_allocation(table4, itgrid_planning_output_json)
    change_table_eco_analyse(table3, grid_planning_output_json, grid_operation_output_json)
    change_table_eco_analyse(table5, itgrid_planning_output_json, itgrid_operation_output_json)
    doc.save('test.docx')

    ###更改图片
    with zipfile.ZipFile('./test.docx') as z:
        z.extractall('./docx/')  # 解压docx文件
    light_image()
    load_image()
    zip_docx('./docx/', 'out.docx')

    ###删除过程文件
    os.remove('./test.docx')
    shutil.rmtree('./docx')



